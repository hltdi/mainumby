#   
#   Mainumby Database for Texts and their Translations.
#   Uses the Object Relational Mapper implementation of SQLAlchemy.
#
########################################################################
#
#   This file is part of the PloGs project
#   for parsing, generation, translation, and computer-assisted
#   human translation.
#
#   Copyright (C) PLoGS <gasser@indiana.edu>
#   
#   This program is free software: you can redistribute it and/or
#   modify it under the terms of the GNU General Public License as
#   published by the Free Software Foundation, either version 3 of
#   the License, or (at your option) any later version.
#   
#   This program is distributed in the hope that it will be useful,
#   but WITHOUT ANY WARRANTY; without even the implied warranty of
#   MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
#   GNU General Public License for more details.
#   
#   You should have received a copy of the GNU General Public License
#   along with this program. If not, see <http://www.gnu.org/licenses/>.
#
# =========================================================================

# 2019.08.12
# -- Created

#from sqlalchemy import create_engine, Column, Integer, String, ForeignKey, DateTime
#from sqlalchemy.ext.declarative import declarative_base
#from sqlalchemy.orm import sessionmaker, relationship, backref
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
import datetime, os
import docx
from .utils import get_time
from .sentence import Document
from .language import Language

# the database class bound to the current app
from . import db

TEXT_DIR = os.path.join(os.path.dirname(__file__), 'texts')
TEXT_EXT = ".txt"
DOCX_EXT = '.docx'

DOMAINS = ["Miscelánea", "Cuentos", "Ciencia", "Gobierno", "Cocina",
           "Historia", "Lenguaje"]

class Human(db.Model):
    """User of the system who is registered and whose feedback is saved."""

    __tablename__ = 'humans'
    __bind_key__ = "text"

    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String)
    fullname = db.Column(db.String)
    email = db.Column(db.String)
    pw_hash = db.Column(db.String)
    # Guarani ability
    level = db.Column(db.Integer)
    creation = db.Column(db.String)

    def __init__(self, username='', email='', password='', name='', level=1, pw_hash=''):
        """name and level are optional. Other fields are required."""
        self.username = username
        self.email = email
        self.level = level
        self.name = name
        self.creation = get_time(True)
        # if password is provided, create the hash, but don't save the password
        if pw_hash:
            self.pw_hash = pw_hash
        elif password:
            self.set_password(password)
        else:
            self.pw_hash = ''

    def __repr__(self):
        return "<Human({}, {})>".format(self.id, self.username)

    def set_password(self, password):
        self.pw_hash = generate_password_hash(password)

    def check_password(self, password):
        res = check_password_hash(self.pw_hash, password)
        return res

    @staticmethod
    def user2human(string):
        """Convert an old User into a Human object, using the string representation of the User."""
        username, pw_hash, email, name, level = string.split(';')
        human = Human(username=username, email=email, name=name, level=level, pw_hash=pw_hash)
        return human

class Text(db.Model):
    """A source-language document stored in a file."""
    
    __tablename__ = 'texts'
    __bind_key__ = "text"

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String)
    title = db.Column(db.String)
    content = db.Column(db.String)
    domain = db.Column(db.String)
    description = db.Column(db.String)

    def __init__(self, name='', domain='Miscelánea', content='', title='',
                 language=None, description=''):
        """Either name or title or both should be specified."""
        if domain not in DOMAINS:
            print("Advertencia: ¡{} no pertenece a la actual lista de dominios!".format(domain))
        self.name = name or title
        self.domain = domain
        self.content = content
        self.title = title or name
        self.language = language
        self.description = description
        self.set_language()

    def __repr__(self):
        return "<Text({}, {}:{})>".format(self.id, self.domain, self.title)

    ### Segmentation
    def segment(self):
        # Let Document do the sentence tokenization work.
        doc = Document(text=self.content, language=self.language)
        for index, sentence in enumerate(doc):
            # Make a TextSeg for each Sentence in the Document
            TextSeg(text=self, content=sentence.original, index=index)

    def set_language(self):
        if not self.language:
            s, t = Language.load_trans('spa', 'grn', train=False)
            self.language = s

    ### Database methods

    @staticmethod
    def delete_all():
        """Remove all Text objects from the database."""
        for t in db.session.query(Text).all():
            db.session.delete(t)

    @staticmethod
    def get_files(name):
        """Text files for this file."""
        all_files = os.listdir(TEXT_DIR)
        return [file for file in all_files if file.split('.')[0] == name]

    @staticmethod
    def get_text_path(name, ext=TEXT_EXT):
        """Path for a file containing a text."""
        return os.path.join(TEXT_DIR, name + ext)

    @staticmethod
    def read(name, domain="Miscelánea", title=''):
        """Read a file in the 'texts' directory and create a Text object
        from its contents."""
        content = ''
        # Get all the files with this name and different extensions
        files = Text.get_files(name)
        # .txt files have priority over .docx
        if any([f.endswith(TEXT_EXT) for f in files]):
            path = Text.get_text_path(name)
            try:
                with open(path, encoding='utf8') as f:
                    content = f.read()
            except IOError:
                print("No se pudo encontrar el archivo {}".format(path))
        elif any([f.endswith(DOCX_EXT) for f in files]):
            content = Text.docx2txt(name)
        if content:
            return Text(name=name, domain=domain, content=content, title=title)
        print("No existe un archivo con nombre {}".format(name))

    @staticmethod
    def docx2txt(name):
        """Extract text from a .docx file."""
        path = Text.get_text_path(name, ext='.docx')
        try:
            doc = docx.Document(path)
            text = [para.text for para in doc.paragraphs]
            # Join paragraphs with something other than 
            text = '\n\n'.join(text)
            return text
        except docx.opc.exceptions.PackageNotFoundError:
            print("No se pudo encontrar el archivo {}".format(path))

class TextSeg(db.Model):
    """A sentence or similar unit within a Text."""
    
    __tablename__ = 'textsegs'
    __bind_key__ = "text"

    id = db.Column(db.Integer, primary_key=True)
    content = db.Column(db.String)
    index = db.Column(db.Integer)
    text_id = db.Column(db.Integer, db.ForeignKey('texts.id'))
    text = db.relationship("Text", backref=db.backref('textsegs', lazy=True),
                           cascade="all, delete")

    def __init__(self, text='', content='', index=0):
        self.text = text
        self.content = content
        self.index = index

    def __repr__(self):
        content = self.content[:25] + '...' if len(self.content) > 25 else self.content
        return "<TextSeg({}, {})>".format(self.id, content)

class Translation(db.Model):
    """A target-language translation of a Text."""
    
    __tablename__ = 'translations'
    __bind_key__ = "text"

    id = db.Column(db.Integer, primary_key=True)
    # The associated Text object
    text_id = db.Column(db.Integer, db.ForeignKey('texts.id'))
    text = db.relationship("Text", backref=db.backref('translations', lazy=True),
                           cascade="all, delete")
    # The associated Human object
    translator_id = db.Column(db.Integer, db.ForeignKey('humans.id'))
    translator = db.relationship("Human", backref=db.backref('humans', lazy=True),
                                 cascade="all, delete")

    def __init__(self, text='', translator=None):
        self.text = text
        self.translator = translator

    def __repr__(self):
        return "<Translation({}, {}, {})>".format(self.id, self.text.name, self.translator.username)

class TraSeg(db.Model):
    """A unit within a translation corresponding to a TextSeg within the associated Text object."""
    
    __tablename__ = 'trasegs'
    __bind_key__ = "text"

    id = db.Column(db.Integer, primary_key=True)
    content = db.Column(db.String)
    index = db.Column(db.Integer)
    translation_id = db.Column(db.Integer, db.ForeignKey('translations.id'))
    translation = db.relationship("Translation", backref=db.backref('trasegs', lazy=True),
                                  cascade="all, delete")

    def __init__(self, content='', translation=None, index=0):
        self.content
        self.translation = translation
        self.index = index

    def __repr__(self):
        content = self.content[:25] + '...' if len(self.content) > 25 else self.content
        return "<TraSeg({}, {})>".format(self.id, content)
